import os
from typing import List
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.pinecone import Pinecone
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.memory import ChatMessageHistory, ConversationBufferMemory
from langchain.docstore.document import Document
import chainlit as cl
import datetime
from main import translate_and_save
from langchain_community.document_loaders import UnstructuredHTMLLoader
from langchain_core.prompts import PromptTemplate
from langchain_openai import OpenAI
from client import llm
from docx import Document
from google.cloud import storage  # GCS client
from langchain.schema import StrOutputParser
from langchain.schema.runnable import Runnable
from langchain.schema.runnable.config import RunnableConfig

import chainlit as cl
from chainlit.types import AskFileResponse

# Initialize GCS client
client = storage.Client(project="sacred-alliance-433217-e3")
bucket_name = "transalation-ndonthi1"  # Replace with your GCS bucket name
bucket = client.bucket(bucket_name)

# index_name = "langchain-demo"
# text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
# embeddings = OpenAIEmbeddings()

# namespaces = set()

welcome_message = """Welcome to the Language Transalation Tool! To get started:
1. Upload a html file
2. Ask to translate the file to a different language
"""


import asyncio
from concurrent.futures import ThreadPoolExecutor

async def upload_to_gcs(file_path, destination_blob_name):
    """Uploads a file to GCS and returns the public URL."""
    loop = asyncio.get_event_loop()
    
    # Run sync upload function in a separate thread
    def sync_upload():
        blob = bucket.blob(destination_blob_name)
        blob.upload_from_filename(file_path)
        url = blob.generate_signed_url(
            version="v4",
            expiration=datetime.timedelta(minutes=15),
            method="GET",
        )
        return url

    # Execute in thread pool to prevent blocking
    return await loop.run_in_executor(None, sync_upload)

def input_text(file):
    loader = UnstructuredHTMLLoader(file.path)
    data = loader.load()
    message_text = data[0].page_content
    return message_text

# Function to save the translation result to a Word document and upload it to GCS
async def save_and_upload_translation(response_text: str, file_name: str) -> str:
    # Create a new Word document
    doc = Document()
    doc.add_heading('Translation Result', level=1)

    # Split the translated text into lines
    lines = response_text.split('\n')

    # Add each line as a separate paragraph
    for line in lines:
        if line.strip():  # Avoid adding empty lines
            doc.add_paragraph(line.strip())

    # Save the document locally
    local_file_path = f'{file_name}.docx'
    doc.save(local_file_path)

    # Upload the document to GCS and return the download URL
    gcs_url = await upload_to_gcs(local_file_path, f'{file_name}.docx')
    return gcs_url

@cl.on_chat_start
async def start():
    files = None
    while files is None:
        files = await cl.AskFileMessage(
            content=welcome_message,
            accept=["text/plain","text/html", "application/pdf"],
            max_size_mb=20,
            timeout=180,
        ).send()

    file = files[0]

    msg = cl.Message(content=f"Processing `{file.name}`...")
    await msg.send()

    # Process the file and get the download link
    gcs_url = await upload_to_gcs(file.path, file.name)

    # Extract text from the file
    message_text = input_text(file)  # Extracts the message_text from the uploaded file
    
    # Prompt to extract language from the user input
    langugae_template = """your task is to extract the name of the language mentioned in the {user_input}. Only respond with the name of the language."""
    langugae_prompt = PromptTemplate(input_variables=["user_input"], template=langugae_template)
    language_chain = langugae_prompt | llm | StrOutputParser()

    # Translation prompt
    translate_template = """Translate the following English text to {target_lang}: {doc_text}.
    There are some Bible references in this text that are enclosed in () so please use the Portuguese King James version to replace them. 
    For the rest, use the best of your translation skills and make sure there are no grammatical mistakes in the translation."""
    translate_prompt = PromptTemplate(input_variables=["target_lang", "doc_text"], template=translate_template)
    
    # Passing extracted message_text to translate_prompt
    runnable = {"target_lang": language_chain, "doc_text": lambda x: message_text} | translate_prompt | llm | StrOutputParser()

    

    # Update the message to include the download link
    msg.content = f"`{file.name}` processed. Uploaded [here]({gcs_url})!"
    await msg.update()
    cl.user_session.set("runnable", runnable)

@cl.on_message
async def main(message: cl.Message):
    runnable = cl.user_session.get("runnable")  # type: ConversationalRetrievalChain
    msg = cl.Message(content="")
    translation_result = ""  # Initialize the translation result variable

    # Stream the response using user input (target language) and extracted text (message_text)
    async for chunk in runnable.astream(
        {"user_input": message.content},  # Passes user input (language) dynamically
        config=RunnableConfig(callbacks=[cl.LangchainCallbackHandler()]),
    ):
        translation_result += chunk  # Accumulate the result
        await msg.stream_token(chunk)

    await msg.send()
    # Once translation is completed, save it to a document and upload to GCS
    file_name = f"translated_file_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
    gcs_url = await save_and_upload_translation(translation_result, file_name)

    # Send the link to the user
    await cl.Message(content=f"Translation completed. Download the file from [here]({gcs_url})!").send()


